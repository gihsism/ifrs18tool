"""Parse financial statement tables from PDF, Word, and image files.

Extraction strategies (tried in order, best result wins):
1. pdfplumber table detection (5 strategies for different layouts)
2. Word-position clustering (handles borderless tables)
3. Layout-preserving text parsing (line-by-line)
4. OCR (for scanned PDFs and images — pytesseract + Pillow)

Multi-page support: tables spanning pages are merged.
"""

import re
import io
import logging
import pandas as pd
import pdfplumber
from collections import Counter
from docx import Document

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Number parsing
# ---------------------------------------------------------------------------

def _clean_number(value) -> float | None:
    """Parse a financial number string.  Handles (parens), currency, thousands sep."""
    if value is None:
        return None
    if not isinstance(value, str):
        try:
            return float(value)
        except (TypeError, ValueError):
            return None
    s = value.strip()
    if not s or s in ("-", "—", "–", "n/a", "N/A", "nil", "Nil", "None", ""):
        return 0.0

    negative = False
    if s.startswith("(") and s.endswith(")"):
        negative = True
        s = s[1:-1]
    elif s.startswith("-") or s.startswith("–") or s.startswith("—"):
        negative = True
        s = s[1:]

    s = re.sub(r"[£$€¥₹\s'\u00a0\u2009]", "", s)
    s = s.replace(",", "").replace(" ", "")

    if s.endswith("-"):
        negative = True
        s = s[:-1]

    if not s:
        return 0.0

    try:
        val = float(s)
        return -val if negative else val
    except ValueError:
        return None


def _is_number_like(value: str) -> bool:
    if not isinstance(value, str):
        return False
    s = value.strip()
    if not s or s in ("-", "—", "–"):
        return False
    result = _clean_number(s)
    return result is not None and result != 0.0


def _is_header_value(value: str) -> bool:
    s = value.strip()
    if not s:
        return False
    if re.match(r"^(FY|CY|PY)?\s*\d{4}$", s, re.IGNORECASE):
        return True
    header_words = [
        "amount", "total", "note", "notes", "current", "prior",
        "year", "period", "restated", "audited", "unaudited",
        "budget", "actual", "forecast", "eur", "usd", "gbp", "chf",
        "rm", "r'000", "'000", "000", "million", "m", "$m", "£m",
        "thousands", "in thousands", "rs", "inr",
    ]
    return s.lower() in header_words or re.match(r"^\d{4}/\d{2,4}$", s)


# ---------------------------------------------------------------------------
# Table scoring
# ---------------------------------------------------------------------------

def _score_table(df: pd.DataFrame) -> float:
    if df.shape[1] < 2 or df.shape[0] < 3:
        return 0

    score = 0
    first_col = df.iloc[:, 0].astype(str)
    text_cells = sum(1 for v in first_col if v.strip() and not _is_number_like(v))
    score += (text_cells / max(len(first_col), 1)) * 25

    num_col_count = 0
    total_num_cells = 0
    for col_idx in range(1, df.shape[1]):
        col_vals = df.iloc[:, col_idx].astype(str)
        num_cells = sum(1 for v in col_vals if _is_number_like(v))
        if num_cells / max(len(col_vals), 1) > 0.2:
            num_col_count += 1
            total_num_cells += num_cells
    score += min(num_col_count, 3) * 10
    score += min(total_num_cells, 20)

    financial_keywords = [
        "revenue", "sales", "cost", "profit", "loss", "income", "expense",
        "tax", "interest", "dividend", "depreciation", "amortis", "ebitda",
        "operating", "gross", "net", "total", "finance", "asset", "liabilit",
        "equity", "cash", "receivable", "payable", "inventory", "provision",
        "impairment", "turnover", "borrowing",
    ]
    keyword_hits = sum(
        1 for v in first_col
        if any(kw in v.lower() for kw in financial_keywords)
    )
    score += (keyword_hits / max(len(first_col), 1)) * 25

    return min(score, 100)


# ---------------------------------------------------------------------------
# Table standardisation
# ---------------------------------------------------------------------------

def _standardise_table(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df = df.replace("", pd.NA).replace("None", pd.NA)
    df = df.dropna(how="all").dropna(axis=1, how="all").fillna("")
    df = df.reset_index(drop=True)

    if df.shape[0] < 2 or df.shape[1] < 2:
        return pd.DataFrame()

    # Detect header row
    header_row_idx = None
    for check_idx in range(min(3, len(df))):
        row_vals = df.iloc[check_idx].astype(str).tolist()
        non_first = row_vals[1:]
        header_like = sum(1 for v in non_first if _is_header_value(v) or not v.strip())
        has_big_number = any(
            _is_number_like(v) and abs(_clean_number(v) or 0) >= 100
            and not re.match(r"^\d{4}$", v.strip())
            for v in non_first
        )
        if header_like >= len(non_first) * 0.5 and not has_big_number:
            header_row_idx = check_idx
            break

    if header_row_idx is not None:
        header = [str(c).strip() for c in df.iloc[header_row_idx]]
        df = df.iloc[header_row_idx + 1:].reset_index(drop=True)
        header[0] = "Account"
        for i in range(1, len(header)):
            if not header[i]:
                header[i] = f"Column {i}"
        df.columns = header
    else:
        df.columns = ["Account"] + [f"Column {i}" for i in range(1, df.shape[1])]

    # Deduplicate column names
    cols = list(df.columns)
    cols[0] = "Account"
    seen = {}
    for i in range(1, len(cols)):
        name = cols[i]
        if name in seen:
            seen[name] += 1
            cols[i] = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
    df.columns = cols

    df["Account"] = df.iloc[:, 0].astype(str).str.strip()

    for i in range(1, len(df.columns)):
        df.iloc[:, i] = df.iloc[:, i].astype(str).apply(_clean_number)

    # Drop empty columns
    cols_to_drop = []
    for i in range(1, len(df.columns)):
        col_series = df.iloc[:, i]
        if col_series.isna().all() or (col_series.fillna(value=0, inplace=False) == 0).all():
            cols_to_drop.append(df.columns[i])
    if cols_to_drop:
        df = df.drop(columns=cols_to_drop)

    df = df[df["Account"].str.strip().str.len() > 0].copy()
    amount_cols = [c for c in df.columns if c != "Account"]
    if amount_cols:
        df = df.dropna(subset=amount_cols, how="all")
    for col in amount_cols:
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

    return df.reset_index(drop=True)


def _final_score(std_df: pd.DataFrame, raw_score: float) -> float:
    """Re-score a standardised table with quality checks."""
    amount_cols = [c for c in std_df.columns if c != "Account"]
    score = raw_score

    if len(amount_cols) >= 2:
        score += 15
    score += min(len(std_df), 30)

    year_cols = sum(
        1 for c in amount_cols if re.match(r"^(FY|CY|PY)?\s*\d{4}$", str(c))
    )
    score += year_cols * 15

    multi_space = sum(1 for v in std_df["Account"] if "  " in str(v))
    score -= (multi_space / max(len(std_df), 1)) * 30

    junk = sum(
        1 for v in std_df["Account"]
        if re.match(r"^(for\s+the|year\s+ended|as\s+at|note\s|account\s)", str(v).lower().strip())
    )
    score -= junk * 10

    name_lengths = [len(str(v).strip()) for v in std_df["Account"]]
    avg_len = sum(name_lengths) / max(len(name_lengths), 1)
    short = sum(1 for l in name_lengths if l < 10)
    if avg_len < 15 and short > len(name_lengths) * 0.3:
        score -= 20

    for col in amount_cols:
        zero_ratio = (std_df[col] == 0).sum() / max(len(std_df), 1)
        if zero_ratio > 0.5:
            score -= 15

    return score


def _dedupe_and_rank(candidates: list[tuple[pd.DataFrame, float]]) -> list[pd.DataFrame]:
    """Standardise, deduplicate, re-score, and rank candidate tables."""
    final = []
    seen = set()

    candidates.sort(key=lambda x: x[1], reverse=True)

    for raw_df, raw_score in candidates:
        try:
            std = _standardise_table(raw_df)
        except Exception:
            continue
        if std.empty or len(std) < 2:
            continue

        sig = (len(std), len(std.columns), tuple(std["Account"].head(5).tolist()))
        if sig in seen:
            continue
        seen.add(sig)

        score = _final_score(std, raw_score)
        final.append((std, score))

    final.sort(key=lambda x: x[1], reverse=True)
    return [df for df, _ in final]


# ---------------------------------------------------------------------------
# PDF: line-by-line text parsing (used by the main extractor and OCR path)
# ---------------------------------------------------------------------------

# Matches financial numbers: 500,000 or (280,000) or -45,000 or 8,500 or 2026
_NUM_TOKEN = re.compile(
    r"\([\s£$€¥]*[\d,]+\.?\d*\)"   # parenthesised: (280,000)
    r"|[\-–—][\s£$€¥]*[\d,]+\.?\d*"  # negative: -45,000
    r"|[\d,]+\.?\d*"                   # plain: 500,000
)


def _parse_financial_line(line: str) -> dict | None:
    """Parse a single text line into account description + number columns.

    Handles both layout-preserved text (numbers separated by spaces/tabs) and
    OCR output (numbers may be separated by single spaces).
    """
    line = line.strip()
    if not line or len(line) < 5:
        return None

    # Find all number tokens in the line
    matches = list(_NUM_TOKEN.finditer(line))
    if not matches:
        return None

    # Identify which matches are "real" financial numbers (not part of text)
    # Walk from the end of the line backwards to find the rightmost cluster of numbers
    real_numbers = []
    for m in reversed(matches):
        token = m.group().strip()
        val = _clean_number(token)
        if val is None:
            continue
        # Skip 4-digit numbers that look like years when they're part of description
        if re.match(r"^\d{4}$", token) and m.start() < len(line) * 0.3:
            continue
        real_numbers.append((m.start(), m.end(), val, token))

    real_numbers.reverse()

    if not real_numbers:
        return None

    # The description is everything before the first real number
    first_num_pos = real_numbers[0][0]
    desc = line[:first_num_pos].strip()
    desc = re.sub(r"[\s\.\-–—:]+$", "", desc)

    if len(desc) < 2:
        return None

    # Extract numbers
    nums = [val for _, _, val, _ in real_numbers]
    if not nums:
        return None

    row = {"Account": desc}
    for i, n in enumerate(nums):
        row[f"Column {i+1}"] = n
    return row


# ---------------------------------------------------------------------------
# OCR: for scanned PDFs and images
# ---------------------------------------------------------------------------

def _ocr_available() -> bool:
    """Check if OCR dependencies are available."""
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _ocr_image(image) -> list[dict]:
    """Run OCR on a PIL Image and parse financial lines.

    Starts with the single best combination (raw image + PSM 6 block layout)
    and only tries alternate preprocessing / PSM modes if the first pass
    yielded fewer than 3 usable rows. This cuts OCR time dramatically on
    normal financial statement images.
    """
    import pytesseract
    from PIL import ImageFilter, ImageOps

    def _run(processed, psm: int) -> list[dict]:
        try:
            text = pytesseract.image_to_string(processed, config=f"--oem 3 --psm {psm}")
        except Exception:
            return []
        rows = []
        for line in text.split("\n"):
            row = _parse_financial_line(line)
            if row:
                rows.append(row)
        return rows

    # Primary pass — the combination that works for ~90% of clean PDFs.
    best = _run(image, 6)
    if len(best) >= 3:
        return best

    # Only fall back if primary failed.
    try:
        sharpened = ImageOps.autocontrast(image.convert("L")).filter(ImageFilter.SHARPEN)
    except Exception:
        sharpened = None

    for processed, psm in [(sharpened, 6), (image, 4), (sharpened, 4)]:
        if processed is None:
            continue
        rows = _run(processed, psm)
        if len(rows) > len(best):
            best = rows
        if len(best) >= 3:
            break

    return best


def _ocr_pdf_pages(file) -> list[tuple[pd.DataFrame, float]]:
    """OCR each page of a PDF as an image."""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return []

    if not _ocr_available():
        return []

    try:
        file.seek(0)
        pdf_bytes = file.read()
        # 200 DPI is ~4x faster than 300 and still accurate enough for OCR of
        # printed financial statements. 300 was overkill.
        images = convert_from_bytes(pdf_bytes, dpi=200)
    except Exception:
        return []

    all_rows = []
    for img in images:
        all_rows.extend(_ocr_image(img))

    if len(all_rows) >= 3:
        df = pd.DataFrame(all_rows).fillna(0)
        score = _score_table(
            pd.DataFrame({str(i): df.iloc[:, i].astype(str) for i in range(df.shape[1])})
        )
        return [(df, max(score, 15))]
    return []


def extract_tables_from_image(file) -> list[pd.DataFrame]:
    """Extract financial tables from an image file (PNG, JPG, etc.)."""
    from PIL import Image

    if not _ocr_available():
        return []

    try:
        img = Image.open(file)
    except Exception:
        return []

    rows = _ocr_image(img)

    if len(rows) < 3:
        return []

    df = pd.DataFrame(rows).fillna(0)
    score = _score_table(
        pd.DataFrame({str(i): df.iloc[:, i].astype(str) for i in range(df.shape[1])})
    )
    return _dedupe_and_rank([(df, max(score, 15))])


# ---------------------------------------------------------------------------
# PDF: multi-page table merging
# ---------------------------------------------------------------------------

def _try_merge_pages(tables: list[pd.DataFrame]) -> list[pd.DataFrame]:
    """Attempt to merge tables from consecutive pages that look like continuations.

    Heuristic: two tables are continuations if they have the same number of columns
    and similar column names.
    """
    if len(tables) <= 1:
        return tables

    merged = [tables[0]]
    for tbl in tables[1:]:
        prev = merged[-1]
        # Same columns? Try to merge
        if (
            list(prev.columns) == list(tbl.columns)
            and len(prev.columns) >= 2
        ):
            # Check that tbl doesn't re-introduce a header
            first_acct = str(tbl.iloc[0].get("Account", "")).lower()
            if first_acct not in ("account", "description", "line item", ""):
                merged[-1] = pd.concat([prev, tbl], ignore_index=True)
                continue
        merged.append(tbl)

    return merged


# ---------------------------------------------------------------------------
# PDF: main entry point
# ---------------------------------------------------------------------------

_PDFPLUMBER_STRATEGIES = [
    {"vertical_strategy": "lines", "horizontal_strategy": "lines"},
    {"vertical_strategy": "text", "horizontal_strategy": "text",
     "snap_tolerance": 10, "join_tolerance": 10,
     "min_words_vertical": 2, "min_words_horizontal": 1},
    {"vertical_strategy": "lines_strict", "horizontal_strategy": "lines_strict"},
    {"vertical_strategy": "lines", "horizontal_strategy": "text",
     "snap_tolerance": 8, "join_tolerance": 8},
]


def extract_tables_from_pdf(file) -> list[pd.DataFrame]:
    """Extract financial tables from a PDF.

    Opens the PDF once and iterates pages once, trying table-extraction
    strategies in order and stopping as soon as the page yields a high-quality
    result (so typical statements finish in one pass, not five). Falls back to
    word-position clustering and finally OCR for scanned pages.
    """
    all_candidates: list[tuple[pd.DataFrame, float]] = []

    try:
        file.seek(0)
    except Exception:
        pass

    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_got_good_table = False

                # Try structured table extraction with early-exit.
                for settings in _PDFPLUMBER_STRATEGIES:
                    try:
                        page_tables = page.extract_tables(table_settings=settings)
                    except Exception:
                        continue
                    for raw_table in page_tables or []:
                        if not raw_table or len(raw_table) < 3:
                            continue
                        df = pd.DataFrame(raw_table).fillna("")
                        score = _score_table(df)
                        if score > 15:
                            all_candidates.append((df, score))
                            if score > 40:
                                page_got_good_table = True
                    if page_got_good_table:
                        break

                if page_got_good_table:
                    continue

                # Fallback 1: word-position clustering for borderless tables.
                wpc = _cluster_page_words(page)
                if wpc is not None:
                    all_candidates.append(wpc)
                    if wpc[1] > 40:
                        continue

                # Fallback 2: line-by-line text parsing.
                try:
                    text = page.extract_text(layout=True, x_density=3, y_density=3) or ""
                except Exception:
                    text = ""
                rows = []
                for line in text.split("\n"):
                    row = _parse_financial_line(line)
                    if row:
                        rows.append(row)
                if len(rows) >= 3:
                    df = pd.DataFrame(rows).fillna(0)
                    score = _score_table(
                        pd.DataFrame({str(i): df.iloc[:, i].astype(str) for i in range(df.shape[1])})
                    )
                    all_candidates.append((df, max(score, 20)))
    except Exception:
        pass

    results = _dedupe_and_rank(all_candidates)

    # OCR only if text extraction yielded nothing — scanned PDF case.
    if not results:
        try:
            file.seek(0)
        except Exception:
            pass
        ocr_candidates = _ocr_pdf_pages(file)
        if ocr_candidates:
            results = _dedupe_and_rank(ocr_candidates)

    if len(results) > 1:
        results = _try_merge_pages(results)

    return results


def _cluster_page_words(page) -> tuple[pd.DataFrame, float] | None:
    """Word-position clustering for a single page (no PDF re-open)."""
    try:
        words = page.extract_words(
            keep_blank_chars=True, x_tolerance=5, y_tolerance=3,
        )
    except Exception:
        return None
    if not words:
        return None

    rows_by_y: dict[int, list] = {}
    for w in words:
        y_key = round(w["top"] / 4) * 4
        rows_by_y.setdefault(y_key, []).append(w)

    sorted_ys = sorted(rows_by_y.keys())
    page_rows = [sorted(rows_by_y[y], key=lambda w: w["x0"]) for y in sorted_ys]
    if len(page_rows) < 4:
        return None

    num_x_positions = []
    for row_words in page_rows:
        for w in row_words:
            if _is_number_like(w["text"]):
                num_x_positions.append(round(w["x1"] / 10) * 10)

    if not num_x_positions:
        return None

    x_counts = Counter(num_x_positions)
    col_rights = sorted([x for x, c in x_counts.items() if c >= 2])
    if not col_rights:
        return None

    merged_cols = [col_rights[0]]
    for x in col_rights[1:]:
        if x - merged_cols[-1] > 30:
            merged_cols.append(x)
        elif x_counts.get(x, 0) > x_counts.get(merged_cols[-1], 0):
            merged_cols[-1] = x

    desc_right = merged_cols[0] - 20 if merged_cols else page.width * 0.5

    table_data = []
    for row_words in page_rows:
        desc_parts = [w["text"] for w in row_words if w["x0"] < desc_right]
        desc = " ".join(desc_parts).strip()

        num_values = [""] * len(merged_cols)
        for w in row_words:
            if w["x0"] >= desc_right - 5:
                best_col = min(
                    range(len(merged_cols)),
                    key=lambda c: abs(w["x1"] - merged_cols[c]),
                )
                existing = num_values[best_col]
                num_values[best_col] = (
                    (existing + " " + w["text"]).strip() if existing else w["text"]
                )

        if desc or any(v for v in num_values):
            table_data.append([desc] + num_values)

    if len(table_data) < 3:
        return None

    df = pd.DataFrame(table_data).fillna("")
    score = _score_table(df)
    if score <= 15:
        return None
    return df, score


# ---------------------------------------------------------------------------
# Word (.docx) parsing
# ---------------------------------------------------------------------------

def extract_tables_from_docx(file) -> list[pd.DataFrame]:
    doc = Document(file)
    tables = []

    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if len(data) < 2:
            continue
        df = pd.DataFrame(data).fillna("")
        score = _score_table(df)
        if score > 10:
            std = _standardise_table(df)
            if not std.empty and len(std) >= 2:
                tables.append((std, score))

    if not tables:
        para_tables = _parse_docx_paragraphs(doc)
        tables.extend(para_tables)

    tables.sort(key=lambda x: x[1], reverse=True)
    return [df for df, _ in tables]


def _parse_docx_paragraphs(doc: Document) -> list[tuple[pd.DataFrame, float]]:
    rows = []
    for para in doc.paragraphs:
        row = _parse_financial_line(para.text)
        if row:
            rows.append(row)

    if len(rows) >= 3:
        df = pd.DataFrame(rows).fillna(0)
        return [(df, 20)]
    return []
